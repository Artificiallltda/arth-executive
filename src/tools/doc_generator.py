# 🛡️ SKILL BLINDADA (14/03/2026) - PADRÃO MANUS AI (GIT PERSISTENCE)
from langchain_core.tools import tool
import os
import re
import uuid
import logging
import unicodedata
import asyncio
import markdown
import base64
from typing import Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from playwright.async_api import async_playwright
from src.config import settings
from pydantic import BaseModel, Field
from src.tools.image_generator import generate_image

logger = logging.getLogger(__name__)

# ─── CONFIGURAÇÃO DE CAMINHOS ABSOLUTOS (MANUS STYLE) ───────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__)) # src/tools
PROJECT_ROOT = os.path.dirname(os.path.dirname(BASE_DIR)) # arth-executive
TEMPLATES_DIR = os.getenv("TEMPLATES_DIR", os.path.join(PROJECT_ROOT, "data", "templates", "docx"))

def _get_docx_template_path(categoria: str = None):
    """Busca template DOCX de forma ABSOLUTA com Fallback (Manus AI)."""
    logger.info(f"🔍 [DOCXGen] Buscando template para categoria: '{categoria}' em {TEMPLATES_DIR}")
    
    if not os.path.exists(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR, exist_ok=True)
        return None

    cat_name = str(categoria or "relatorio").lower().replace("template_", "").replace(".docx", "")
    template_path = os.path.join(TEMPLATES_DIR, f"template_{cat_name}.docx")

    if not os.path.exists(template_path):
        logger.warning(f"⚠️ [DOCXGen] Template '{template_path}' não encontrado. Usando fallback.")
        template_path = os.path.join(TEMPLATES_DIR, "template_relatorio.docx")

    if not os.path.exists(template_path):
        logger.error(f"❌ [DOCXGen] Nenhum template DOCX em {template_path}.")
        return None
    
    logger.info(f"✅ [DOCXGen] Usando template absoluto: {template_path}")
    return template_path

def _safe_filename(title: str, max_len: int = 50) -> str:
    normalized = unicodedata.normalize('NFKD', title)
    ascii_only = normalized.encode('ascii', errors='ignore').decode('ascii')
    safe = re.sub(r'[^\w\s-]', '', ascii_only).strip().lower().replace(' ', '-')
    return (safe or 'documento')[:max_len]

# ─── TEMPLATE HTML DE LUXO ──────────────────────────────────────────────────
HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;600;700&display=swap');
        body {{ font-family: 'Inter', sans-serif; -webkit-print-color-adjust: exact; }}
        .markdown-content h1 {{ @apply text-3xl font-bold text-slate-900 mb-6 border-b pb-4; }}
        .markdown-content h2 {{ @apply text-2xl font-semibold text-blue-800 mt-10 mb-4 flex items-center; }}
        .markdown-content h2::before {{ content: ''; @apply w-1.5 h-6 bg-blue-600 mr-3 rounded-full; }}
        .markdown-content p {{ @apply text-slate-700 leading-relaxed mb-4 text-justify; }}
        .markdown-content table {{ @apply w-full border-collapse border border-slate-200 my-8 shadow-sm rounded-lg overflow-hidden; }}
        .markdown-content th {{ @apply bg-slate-100 p-4 text-left font-semibold text-slate-900 border-b border-slate-200; }}
        .markdown-content td {{ @apply p-4 border-b border-slate-100 text-slate-700; }}
    </style>
</head>
<body class="bg-white p-0">
    <div class="flex justify-between items-center mb-12 border-b-2 border-slate-900 pb-4">
        <div class="text-2xl font-bold tracking-tighter text-slate-900">ARTH <span class="text-blue-600">EXECUTIVE</span></div>
        <div class="text-right text-[10px] text-slate-400 uppercase tracking-widest font-semibold">CONFIDENCIAL</div>
    </div>
    {cover_image_html}
    <h1 class="text-5xl font-black text-slate-950 uppercase mt-8">{title}</h1>
    <div class="text-slate-500 text-sm mb-12">Gerado em {date} &middot; Orion Elite v3</div>
    <div class="markdown-content">{content_html}</div>
</body>
</html>
"""

class PdfSchema(BaseModel):
    title: str = Field(..., description="Título do PDF.")
    content: str = Field(..., description="Conteúdo em markdown.")
    include_image: bool = Field(default=False, description="Gerar capa via AI.")

class DocxSchema(BaseModel):
    title: str = Field(..., description="Título do documento.")
    content: str = Field(..., description="Conteúdo em markdown.")
    filename: Optional[str] = Field(None, description="Nome do arquivo.")
    template_name: Optional[str] = Field(None, description="Categoria do template.")

@tool(args_schema=PdfSchema)
async def generate_pdf(title: str, content: str, include_image: bool = False) -> str:
    """Gera PDF de Alta Fidelidade."""
    uid, safe_name = uuid.uuid4().hex[:6], _safe_filename(title)
    filename = f"Exec-{safe_name}-{uid}.pdf"
    filepath = os.path.join(settings.DATA_OUTPUTS_PATH, filename)
    
    cover_image_html = ""
    if include_image:
        try:
            img_res = await generate_image.ainvoke({"prompt": f"Corporate cover for {title}", "orientation": "horizontal"})
            match = re.search(r'SEND_FILE:(img-[^\.]+)', img_res)
            if match:
                img_path = os.path.join(settings.DATA_OUTPUTS_PATH, f"{match.group(1)}.png")
                if os.path.exists(img_path):
                    with open(img_path, "rb") as f:
                        encoded = base64.b64encode(f.read()).decode()
                        cover_image_html = f'<img src="data:image/png;base64,{encoded}" class="w-full h-64 object-cover rounded-xl mb-8">'
        except: pass

    try:
        content_html = markdown.markdown(content, extensions=['tables', 'fenced_code', 'nl2br'])
        full_html = HTML_TEMPLATE.format(title=title.upper(), content_html=content_html, date=datetime.now().strftime("%d/%m/%Y"), cover_image_html=cover_image_html)
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            await page.set_content(full_html)
            await page.wait_for_timeout(2000)
            await page.pdf(path=filepath, format="A4", print_background=True)
            await browser.close()
        return f"Relatório PDF gerado: <SEND_FILE:{filename}>"
    except Exception as e: return f"Erro PDF: {str(e)}"

@tool(args_schema=DocxSchema)
async def generate_docx(title: str, content: str, filename: Optional[str] = None, template_name: Optional[str] = None) -> str:
    """Gera DOCX Profissional com suporte a TEMPLATES PERSISTENTES."""
    fn = filename or f"Doc-{_safe_filename(title)}-{uuid.uuid4().hex[:6]}.docx"
    filepath = os.path.join(settings.DATA_OUTPUTS_PATH, fn)
    def _build():
        tp = _get_docx_template_path(template_name)
        doc = Document(tp) if tp else Document()
        doc.add_heading(title, 0)
        for line in content.split('\n'):
            if line.startswith('# '): doc.add_heading(line[2:], level=1)
            elif line.startswith('## '): doc.add_heading(line[3:], level=2)
            elif line.startswith('- '): doc.add_paragraph(line[2:], style='List Bullet')
            elif line.strip(): doc.add_paragraph(line)
        doc.save(filepath)
    await asyncio.to_thread(_build)
    return f"Documento Word gerado: <SEND_FILE:{fn}>"
