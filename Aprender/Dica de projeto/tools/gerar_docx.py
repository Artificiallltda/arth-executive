"""
GERADOR DE DOCX PROFISSIONAL
Usa: python-docx com design e formatação avançados
"""

import os
import uuid
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def _set_cell_background(cell, hex_color: str):
    """Define cor de fundo de célula de tabela."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def gerar_docx(titulo: str, secoes: list, autor: str = "Sistema IA") -> str:
    """
    Gera DOCX profissional.
    
    secoes: lista de dicts com:
      - titulo: str
      - conteudo: str
      - tipo: "texto" | "lista" | "tabela"
      - dados: list (para tabelas) | list (para listas)
    """
    doc = Document()
    
    # ── Configurar página ──
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    
    # ── Capa / Título Principal ──
    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo_p.add_run(titulo.upper())
    run_titulo.font.size = Pt(26)
    run_titulo.font.bold = True
    run_titulo.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    run_titulo.font.name = "Segoe UI"
    
    # Linha decorativa
    linha = doc.add_paragraph()
    linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_linha = linha.add_run("─" * 50)
    run_linha.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
    run_linha.font.size = Pt(10)
    
    # Autor e data
    from datetime import datetime
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta.add_run(f"{autor} | {datetime.now().strftime('%d/%m/%Y')}")
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run_meta.font.italic = True
    
    doc.add_paragraph()
    
    # ── Seções ──
    for secao in secoes:
        tipo = secao.get("tipo", "texto")
        titulo_secao = secao.get("titulo", "")
        conteudo = secao.get("conteudo", "")
        
        # Título da seção
        if titulo_secao:
            h = doc.add_heading(titulo_secao, level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
                run.font.name = "Segoe UI"
                run.font.size = Pt(14)
        
        if tipo == "texto":
            p = doc.add_paragraph(conteudo)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.name = "Calibri"
        
        elif tipo == "lista":
            itens = secao.get("dados", [])
            for item in itens:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(item)
                run.font.size = Pt(11)
                run.font.name = "Calibri"
        
        elif tipo == "tabela":
            dados_tabela = secao.get("dados", [])
            if dados_tabela:
                cabecalhos = dados_tabela[0]
                linhas = dados_tabela[1:]
                
                tabela = doc.add_table(rows=1, cols=len(cabecalhos))
                tabela.style = "Table Grid"
                
                # Cabeçalho
                row_cab = tabela.rows[0]
                for i, cab in enumerate(cabecalhos):
                    cell = row_cab.cells[i]
                    cell.text = cab
                    _set_cell_background(cell, "1E3A8A")
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Dados
                for j, linha in enumerate(linhas):
                    row = tabela.add_row()
                    for i, valor in enumerate(linha):
                        cell = row.cells[i]
                        cell.text = str(valor)
                        if j % 2 == 0:
                            _set_cell_background(cell, "DBEAFE")
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(10)
        
        doc.add_paragraph()
    
    # ── Rodapé ──
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = footer_p.add_run(f"{titulo} | Gerado por Sistema IA")
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    # Salvar
    os.makedirs("outputs", exist_ok=True)
    uid = str(uuid.uuid4())[:8]
    caminho = f"outputs/{titulo.replace(' ', '_')}_{uid}.docx"
    doc.save(caminho)
    print(f"✅ DOCX gerado: {caminho}")
    return caminho
