"""
CRIADOR DE TEMPLATES DOCX PROFISSIONAIS
Categorias: Financeiro, Marketing, Relatório, Escola, Ideias, Corporativo, Saúde, Tech
Uso: python criar_templates_docx.py
Resultado: arquivos .docx na pasta templates/docx/
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.makedirs("templates/docx", exist_ok=True)

# ── Paletas de cores por categoria ──
PALETAS_DOCX = {
    "financeiro": {
        "titulo":    RGBColor(0x0A, 0x1F, 0x44),
        "subtitulo": RGBColor(0x00, 0x8B, 0x8B),
        "acento":    RGBColor(0xF0, 0xB4, 0x29),
        "texto":     RGBColor(0x1E, 0x29, 0x3B),
        "cab_hex":   "0A1F44",
        "linha_hex": "008B8B",
    },
    "marketing": {
        "titulo":    RGBColor(0x6D, 0x28, 0xD9),
        "subtitulo": RGBColor(0xFF, 0x00, 0x80),
        "acento":    RGBColor(0xFF, 0x6B, 0x35),
        "texto":     RGBColor(0x1E, 0x29, 0x3B),
        "cab_hex":   "6D28D9",
        "linha_hex": "FF0080",
    },
    "relatorio": {
        "titulo":    RGBColor(0x1E, 0x3A, 0x8A),
        "subtitulo": RGBColor(0x3B, 0x82, 0xF6),
        "acento":    RGBColor(0x06, 0xB6, 0xD4),
        "texto":     RGBColor(0x1E, 0x29, 0x3B),
        "cab_hex":   "1E3A8A",
        "linha_hex": "3B82F6",
    },
    "escola": {
        "titulo":    RGBColor(0x05, 0x96, 0x68),
        "subtitulo": RGBColor(0xF5, 0x9E, 0x0B),
        "acento":    RGBColor(0x10, 0xB9, 0x81),
        "texto":     RGBColor(0x06, 0x4E, 0x3B),
        "cab_hex":   "059668",
        "linha_hex": "F59E0B",
    },
    "ideias": {
        "titulo":    RGBColor(0xD9, 0x77, 0x06),
        "subtitulo": RGBColor(0x7C, 0x3A, 0xED),
        "acento":    RGBColor(0xF5, 0x9E, 0x0B),
        "texto":     RGBColor(0x1C, 0x19, 0x17),
        "cab_hex":   "D97706",
        "linha_hex": "7C3AED",
    },
    "corporativo": {
        "titulo":    RGBColor(0x1E, 0x3A, 0x5F),
        "subtitulo": RGBColor(0xC0, 0x39, 0x2B),
        "acento":    RGBColor(0xF3, 0x9C, 0x12),
        "texto":     RGBColor(0x1E, 0x29, 0x3B),
        "cab_hex":   "1E3A5F",
        "linha_hex": "C0392B",
    },
    "saude": {
        "titulo":    RGBColor(0x05, 0x7A, 0x55),
        "subtitulo": RGBColor(0x06, 0x95, 0xDD),
        "acento":    RGBColor(0x10, 0xB9, 0x81),
        "texto":     RGBColor(0x06, 0x4E, 0x3B),
        "cab_hex":   "057A55",
        "linha_hex": "0695DD",
    },
    "tech": {
        "titulo":    RGBColor(0x0D, 0x11, 0x17),
        "subtitulo": RGBColor(0x00, 0xC8, 0xA0),
        "acento":    RGBColor(0x7C, 0x3A, 0xED),
        "texto":     RGBColor(0x1E, 0x29, 0x3B),
        "cab_hex":   "0D1117",
        "linha_hex": "00C8A0",
    },
}

def _set_cell_bg(cell, hex_color: str):
    """Define cor de fundo de célula."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _add_linha_colorida(doc, hex_color: str, espessura: int = 4):
    """Adiciona linha horizontal colorida."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(espessura * 4))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def _configurar_cabecalho(doc, titulo_doc: str, paleta: dict):
    """Configura cabeçalho com fundo colorido."""
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    
    # Limpar parágrafo padrão
    for p in header.paragraphs:
        p.clear()
    
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"  {titulo_doc.upper()}")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "Segoe UI"
    
    # Fundo colorido no cabeçalho via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), paleta["cab_hex"])
    pPr.append(shd)

def _configurar_rodape(doc, nome_template: str):
    """Configura rodapé com número de página."""
    footer = doc.sections[0].footer
    for p in footer.paragraphs:
        p.clear()
    
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Template {nome_template.capitalize()} | Gerado por Sistema IA | Página ")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    # Número de página automático
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run2 = p.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

def criar_template_docx(nome: str, paleta: dict, descricao: str):
    """Cria um template DOCX completo com estrutura profissional."""
    doc = Document()

    # ── Configurar página ──
    section = doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    # ── Cabeçalho e rodapé ──
    _configurar_cabecalho(doc, f"Template {nome}", paleta)
    _configurar_rodape(doc, nome)

    # ── Título principal ──
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_titulo.add_run(f"TÍTULO DO DOCUMENTO")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = paleta["titulo"]
    run.font.name = "Segoe UI"

    # ── Linha decorativa ──
    _add_linha_colorida(doc, paleta["linha_hex"], espessura=3)

    # ── Subtítulo / Descrição ──
    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run(f"Subtítulo ou descrição do documento | {descricao}")
    run_sub.font.size = Pt(12)
    run_sub.font.italic = True
    run_sub.font.color.rgb = paleta["subtitulo"]
    run_sub.font.name = "Segoe UI"

    doc.add_paragraph()

    # ── Seção 1: Introdução ──
    h1 = doc.add_heading("1. Introdução", level=2)
    for run in h1.runs:
        run.font.color.rgb = paleta["titulo"]
        run.font.name = "Segoe UI"
        run.font.size = Pt(14)

    p1 = doc.add_paragraph(
        "Insira aqui o texto introdutório do documento. Este template foi criado para uso profissional "
        "e pode ser personalizado conforme necessário. O conteúdo deve ser claro, objetivo e bem estruturado."
    )
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p1.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = paleta["texto"]

    doc.add_paragraph()

    # ── Seção 2: Desenvolvimento ──
    h2 = doc.add_heading("2. Desenvolvimento", level=2)
    for run in h2.runs:
        run.font.color.rgb = paleta["titulo"]
        run.font.name = "Segoe UI"
        run.font.size = Pt(14)

    p2 = doc.add_paragraph(
        "Desenvolva aqui o conteúdo principal do documento. Utilize subtítulos para organizar "
        "as informações e facilitar a leitura. Mantenha parágrafos curtos e objetivos."
    )
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p2.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = paleta["texto"]

    doc.add_paragraph()

    # ── Subtítulo 2.1 ──
    h2_1 = doc.add_heading("2.1 Subtópico", level=3)
    for run in h2_1.runs:
        run.font.color.rgb = paleta["subtitulo"]
        run.font.name = "Segoe UI"
        run.font.size = Pt(12)

    # ── Lista de exemplo ──
    itens = [
        "Primeiro item da lista com informação relevante",
        "Segundo item com dado importante",
        "Terceiro ponto de destaque",
        "Quarto elemento da lista",
    ]
    for item in itens:
        p_item = doc.add_paragraph(style="List Bullet")
        run_item = p_item.add_run(item)
        run_item.font.size = Pt(11)
        run_item.font.name = "Calibri"
        run_item.font.color.rgb = paleta["texto"]

    doc.add_paragraph()

    # ── Seção 3: Tabela de dados ──
    h3 = doc.add_heading("3. Dados e Análises", level=2)
    for run in h3.runs:
        run.font.color.rgb = paleta["titulo"]
        run.font.name = "Segoe UI"
        run.font.size = Pt(14)

    # Tabela
    tabela = doc.add_table(rows=1, cols=4)
    tabela.style = "Table Grid"

    cabecalhos = ["Item", "Descrição", "Valor", "Status"]
    row_cab = tabela.rows[0]
    for i, cab in enumerate(cabecalhos):
        cell = row_cab.cells[i]
        cell.text = cab
        _set_cell_bg(cell, paleta["cab_hex"])
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            run.font.name = "Segoe UI"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    dados_exemplo = [
        ["Item 01", "Descrição do primeiro item", "R$ 1.000,00", "Ativo"],
        ["Item 02", "Descrição do segundo item", "R$ 2.500,00", "Pendente"],
        ["Item 03", "Descrição do terceiro item", "R$ 800,00", "Concluído"],
    ]
    for j, linha in enumerate(dados_exemplo):
        row = tabela.add_row()
        fill = "DBEAFE" if j % 2 == 0 else "FFFFFF"
        for i, valor in enumerate(linha):
            cell = row.cells[i]
            cell.text = valor
            _set_cell_bg(cell, fill)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"

    doc.add_paragraph()

    # ── Seção 4: Conclusão ──
    h4 = doc.add_heading("4. Conclusão", level=2)
    for run in h4.runs:
        run.font.color.rgb = paleta["titulo"]
        run.font.name = "Segoe UI"
        run.font.size = Pt(14)

    p_conc = doc.add_paragraph(
        "Insira aqui as considerações finais do documento. Resuma os principais pontos abordados "
        "e apresente as recomendações ou próximos passos quando aplicável."
    )
    p_conc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p_conc.runs:
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = paleta["texto"]

    # ── Linha final ──
    _add_linha_colorida(doc, paleta["linha_hex"], espessura=2)
    p_final = doc.add_paragraph()
    run_final = p_final.add_run(f"Documento gerado pelo Sistema IA | Template: {nome.capitalize()}")
    run_final.font.size = Pt(9)
    run_final.font.italic = True
    run_final.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caminho = f"templates/docx/template_{nome}.docx"
    doc.save(caminho)
    print(f"✅ Template DOCX criado: {caminho}")
    return caminho


# ── Criar todos os templates ──
if __name__ == "__main__":
    configs = [
        ("financeiro",  PALETAS_DOCX["financeiro"],  "Relatórios financeiros, balanços e análises"),
        ("marketing",   PALETAS_DOCX["marketing"],   "Campanhas, estratégias e planos de marketing"),
        ("relatorio",   PALETAS_DOCX["relatorio"],   "Relatórios gerais e documentos técnicos"),
        ("escola",      PALETAS_DOCX["escola"],      "Trabalhos acadêmicos, planos de aula e atividades"),
        ("ideias",      PALETAS_DOCX["ideias"],      "Brainstorming, projetos criativos e inovação"),
        ("corporativo", PALETAS_DOCX["corporativo"], "Documentos corporativos e institucionais"),
        ("saude",       PALETAS_DOCX["saude"],       "Relatórios médicos e documentos de saúde"),
        ("tech",        PALETAS_DOCX["tech"],        "Documentação técnica e projetos de tecnologia"),
    ]

    for nome, paleta, descricao in configs:
        criar_template_docx(nome, paleta, descricao)

    print("\n✅ Todos os templates DOCX foram criados em templates/docx/")
