"""
GERADOR DE DOCX COM IMAGENS CONTEXTUAIS
Fluxo completo: Gemini cria prompt → DALL-E gera imagem → imagem entra no documento
"""

import os
import uuid
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from gerar_imagem_contextual import gerar_imagem_contextual

os.makedirs("outputs", exist_ok=True)

# ── Paletas por categoria ──
PALETAS_DOCX = {
    "financeiro":   {"titulo": RGBColor(0x0A,0x1F,0x44), "subtitulo": RGBColor(0x00,0x8B,0x8B), "acento": RGBColor(0xF0,0xB4,0x29), "texto": RGBColor(0x1E,0x29,0x3B), "cab_hex": "0A1F44", "linha_hex": "008B8B"},
    "marketing":    {"titulo": RGBColor(0x6D,0x28,0xD9), "subtitulo": RGBColor(0xFF,0x00,0x80), "acento": RGBColor(0xFF,0x6B,0x35), "texto": RGBColor(0x1E,0x29,0x3B), "cab_hex": "6D28D9", "linha_hex": "FF0080"},
    "relatorio":    {"titulo": RGBColor(0x1E,0x3A,0x8A), "subtitulo": RGBColor(0x3B,0x82,0xF6), "acento": RGBColor(0x06,0xB6,0xD4), "texto": RGBColor(0x1E,0x29,0x3B), "cab_hex": "1E3A8A", "linha_hex": "3B82F6"},
    "escola":       {"titulo": RGBColor(0x05,0x96,0x68), "subtitulo": RGBColor(0xF5,0x9E,0x0B), "acento": RGBColor(0x10,0xB9,0x81), "texto": RGBColor(0x06,0x4E,0x3B), "cab_hex": "059668", "linha_hex": "F59E0B"},
    "ideias":       {"titulo": RGBColor(0xD9,0x77,0x06), "subtitulo": RGBColor(0x7C,0x3A,0xED), "acento": RGBColor(0xF5,0x9E,0x0B), "texto": RGBColor(0x1C,0x19,0x17), "cab_hex": "D97706", "linha_hex": "7C3AED"},
    "corporativo":  {"titulo": RGBColor(0x1E,0x3A,0x5F), "subtitulo": RGBColor(0xC0,0x39,0x2B), "acento": RGBColor(0xF3,0x9C,0x12), "texto": RGBColor(0x1E,0x29,0x3B), "cab_hex": "1E3A5F", "linha_hex": "C0392B"},
    "saude":        {"titulo": RGBColor(0x05,0x7A,0x55), "subtitulo": RGBColor(0x06,0x95,0xDD), "acento": RGBColor(0x10,0xB9,0x81), "texto": RGBColor(0x06,0x4E,0x3B), "cab_hex": "057A55", "linha_hex": "0695DD"},
    "tech":         {"titulo": RGBColor(0x0D,0x11,0x17), "subtitulo": RGBColor(0x00,0xC8,0xA0), "acento": RGBColor(0x7C,0x3A,0xED), "texto": RGBColor(0x1E,0x29,0x3B), "cab_hex": "0D1117", "linha_hex": "00C8A0"},
}

def _set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def _add_linha(doc, hex_color, espessura=3):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), str(espessura * 4))
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom); pPr.append(pBdr)
    return p


def gerar_docx_com_imagens(
    titulo: str,
    secoes: list,
    categoria: str = "relatorio",
    autor: str = "Sistema IA",
    gerar_imagens: bool = True,
    imagem_capa: bool = True,
    caminho_saida: str = None
) -> str:
    """
    Gera DOCX profissional com imagens contextuais geradas por IA.

    Args:
        titulo: Título do documento
        secoes: Lista de dicts:
            [{"titulo": "...", "conteudo": "...", "tipo": "texto|lista|tabela", "imagem": True/False}]
        categoria: Categoria do template (define cores e estilo visual)
        autor: Nome do autor
        gerar_imagens: Se True, gera imagens com Gemini + DALL-E
        imagem_capa: Se True, gera imagem de destaque no início do documento
        caminho_saida: Caminho para salvar o arquivo

    Returns:
        Caminho do arquivo DOCX gerado
    """
    paleta = PALETAS_DOCX.get(categoria, PALETAS_DOCX["relatorio"])

    doc = Document()

    # ── Configurar página ──
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    # ── Cabeçalho colorido ──
    header = doc.sections[0].header
    for p in header.paragraphs:
        p.clear()
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_h = p_header.add_run(f"  {titulo.upper()}")
    run_h.font.bold = True; run_h.font.size = Pt(10)
    run_h.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run_h.font.name = "Segoe UI"
    pPr = p_header._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), paleta["cab_hex"])
    pPr.append(shd)

    # ── Rodapé ──
    footer = doc.sections[0].footer
    for p in footer.paragraphs:
        p.clear()
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = p_footer.add_run(f"{titulo} | {autor} | Página ")
    run_f.font.size = Pt(9); run_f.font.color.rgb = RGBColor(0x94,0xA3,0xB8)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run_pg = p_footer.add_run()
    run_pg._r.append(fldChar1); run_pg._r.append(instrText); run_pg._r.append(fldChar2)
    run_pg.font.size = Pt(9); run_pg.font.color.rgb = RGBColor(0x94,0xA3,0xB8)

    # ── Imagem de capa (banner de destaque) ──
    if gerar_imagens and imagem_capa:
        try:
            print(f"🎨 Gerando imagem de capa para: '{titulo}'")
            img_capa = gerar_imagem_contextual(
                tema=titulo,
                categoria=categoria,
                tipo_imagem="banner de cabeçalho de documento profissional, wide horizontal",
                tamanho="horizontal"
            )
            doc.add_picture(img_capa, width=Inches(6.0))
            ultimo_para = doc.paragraphs[-1]
            ultimo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"⚠️ Imagem de capa não gerada: {e}")

    # ── Título principal ──
    p_titulo = doc.add_paragraph()
    run_t = p_titulo.add_run(titulo.upper())
    run_t.font.size = Pt(26); run_t.font.bold = True
    run_t.font.color.rgb = paleta["titulo"]; run_t.font.name = "Segoe UI"

    _add_linha(doc, paleta["linha_hex"])

    from datetime import datetime
    p_meta = doc.add_paragraph()
    run_meta = p_meta.add_run(f"{autor}  |  {datetime.now().strftime('%d/%m/%Y')}  |  Categoria: {categoria.capitalize()}")
    run_meta.font.size = Pt(10); run_meta.font.italic = True
    run_meta.font.color.rgb = paleta["subtitulo"]

    doc.add_paragraph()

    # ── Seções de conteúdo ──
    for i, secao in enumerate(secoes, 1):
        tipo = secao.get("tipo", "texto")
        titulo_secao = secao.get("titulo", "")
        conteudo = secao.get("conteudo", "")
        gerar_img_secao = secao.get("imagem", False)

        # Título da seção
        if titulo_secao:
            h = doc.add_heading(f"{i}. {titulo_secao}", level=2)
            for run in h.runs:
                run.font.color.rgb = paleta["titulo"]
                run.font.name = "Segoe UI"; run.font.size = Pt(13)

        # Imagem contextual para a seção (se solicitado)
        if gerar_imagens and gerar_img_secao and titulo_secao:
            try:
                print(f"🎨 Gerando imagem para seção: '{titulo_secao}'")
                img_secao = gerar_imagem_contextual(
                    tema=f"{titulo_secao} — {titulo}",
                    categoria=categoria,
                    tipo_imagem="ilustração de seção de documento",
                    tamanho="horizontal",
                    qualidade="standard"  # standard para economizar créditos
                )
                doc.add_picture(img_secao, width=Inches(4.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            except Exception as e:
                print(f"⚠️ Imagem da seção não gerada: {e}")

        # Conteúdo
        if tipo == "texto":
            p = doc.add_paragraph(str(conteudo))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.size = Pt(11); run.font.name = "Calibri"
                run.font.color.rgb = paleta["texto"]

        elif tipo == "lista":
            itens = secao.get("dados", [])
            if not itens and isinstance(conteudo, list):
                itens = conteudo
            for item in itens:
                p_item = doc.add_paragraph(style="List Bullet")
                run_item = p_item.add_run(str(item))
                run_item.font.size = Pt(11); run_item.font.name = "Calibri"

        elif tipo == "tabela":
            dados_tabela = secao.get("dados", [])
            if dados_tabela and len(dados_tabela) > 1:
                cabecalhos = dados_tabela[0]
                linhas = dados_tabela[1:]
                tabela = doc.add_table(rows=1, cols=len(cabecalhos))
                tabela.style = "Table Grid"
                row_cab = tabela.rows[0]
                for j, cab in enumerate(cabecalhos):
                    cell = row_cab.cells[j]; cell.text = cab
                    _set_cell_bg(cell, paleta["cab_hex"])
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                        run.font.size = Pt(10); run.font.name = "Segoe UI"
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for k, linha in enumerate(linhas):
                    row = tabela.add_row()
                    fill = "DBEAFE" if k % 2 == 0 else "FFFFFF"
                    for j, valor in enumerate(linha):
                        cell = row.cells[j]; cell.text = str(valor)
                        _set_cell_bg(cell, fill)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(10); run.font.name = "Calibri"

        doc.add_paragraph()

    # ── Linha final ──
    _add_linha(doc, paleta["linha_hex"], espessura=2)
    p_final = doc.add_paragraph()
    run_final = p_final.add_run(f"Documento gerado por Sistema IA com Google Gemini | {titulo}")
    run_final.font.size = Pt(9); run_final.font.italic = True
    run_final.font.color.rgb = RGBColor(0x94,0xA3,0xB8)
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salvar
    if not caminho_saida:
        uid = str(uuid.uuid4())[:8]
        caminho_saida = f"outputs/{titulo.replace(' ', '_')}_{uid}.docx"

    doc.save(caminho_saida)
    print(f"\n✅ DOCX com imagens gerado: {caminho_saida}")
    return caminho_saida


# ── Exemplo de uso ──
if __name__ == "__main__":
    secoes = [
        {
            "titulo": "Visão Geral",
            "conteudo": "Este relatório apresenta os principais resultados financeiros do primeiro trimestre de 2026.",
            "tipo": "texto",
            "imagem": True   # ← gera imagem contextual para esta seção
        },
        {
            "titulo": "Indicadores Principais",
            "conteudo": ["Receita: R$ 2,4M (+35%)", "Margem: 22%", "CAC: R$ 85", "LTV: R$ 1.200"],
            "tipo": "lista",
            "imagem": False
        },
        {
            "titulo": "Tabela de Resultados",
            "tipo": "tabela",
            "imagem": False,
            "dados": [
                ["Produto", "Receita", "Crescimento"],
                ["Produto A", "R$ 1.200.000", "+42%"],
                ["Produto B", "R$ 800.000", "+28%"],
                ["Produto C", "R$ 400.000", "+15%"],
            ]
        },
    ]

    gerar_docx_com_imagens(
        titulo="Relatório Financeiro Q1 2026",
        secoes=secoes,
        categoria="financeiro",
        autor="Gean",
        gerar_imagens=True,
        imagem_capa=True
    )
