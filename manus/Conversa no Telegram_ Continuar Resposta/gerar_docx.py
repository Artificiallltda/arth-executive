from docx import Document
from docx.shared import Inches

def gerar_docx_exemplo(nome_arquivo):
    document = Document()

    document.add_heading('Relatório de Exemplo', level=1)

    document.add_paragraph('Este é um parágrafo de exemplo para demonstrar a geração de documentos DOCX com python-docx.')

    document.add_heading('Seção 1: Dados Importantes', level=2)
    document.add_paragraph('Aqui poderiam vir dados dinâmicos, como resultados de análises ou informações de um banco de dados.')

    # Adicionar uma tabela
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item'
    hdr_cells[1].text = 'Quantidade'
    hdr_cells[2].text = 'Valor'

    # Exemplo de dados para a tabela
    dados = [
        ('Produto A', 10, 150.00),
        ('Produto B', 5, 200.50),
        ('Produto C', 12, 75.25),
    ]

    for item, qtd, valor in dados:
        row_cells = table.add_row().cells
        row_cells[0].text = str(item)
        row_cells[1].text = str(qtd)
        row_cells[2].text = f'R$ {valor:.2f}'

    document.add_page_break()

    document.add_heading('Seção 2: Conclusão', level=2)
    document.add_paragraph('O relatório foi gerado com sucesso utilizando a biblioteca python-docx.')

    document.save(nome_arquivo)
    print(f'Documento {nome_arquivo} gerado com sucesso.')

if __name__ == '__main__':
    gerar_docx_exemplo('relatorio_exemplo.docx')
